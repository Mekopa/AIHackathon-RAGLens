import logging
import os
import tempfile
import glob
import subprocess
import re
from langdetect import detect, LangDetectException
from .base import TextExtractor

logger = logging.getLogger(__name__)

# Define supported languages and their tesseract codes
LANGUAGE_CODES = {
    'en': 'eng',     # English
    'lt': 'lit',     # Lithuanian
    'tr': 'tur'      # Turkish
}

def detect_language(text):
    """
    Detect language of the given text
    
    Args:
        text: Input text
        
    Returns:
        str: Tesseract language code or 'eng' as fallback
    """
    if not text or len(text.strip()) < 20:
        return 'eng'  # Default to English for very short texts
    
    # First, check for Lithuanian-specific filename patterns
    # This helps with DOC files where character encoding might be problematic
    if any(word in text.lower() for word in ['teism', 'lietuv', 'valstyb', 'teisė', 'įstatymas', 'nutartis']):
        logger.info("Found Lithuanian keywords in text, assuming Lithuanian document")
        return 'lit+eng'

    # Check for specific character markers for Lithuanian and Turkish
    # Check early before other detection methods
    if re.search(r'[\u0104\u0105\u010C\u010D\u0116\u0117\u012E\u012F\u0160\u0161\u016A\u016B\u017D\u017E]', text[:10000]):
        logger.info("Found Lithuanian characters, using Lithuanian language")
        return 'lit+eng'
    elif re.search(r'[\u00C7\u00E7\u011E\u011F\u0130\u0131\u015E\u015F\u00D6\u00F6\u00DC\u00FC]', text[:10000]):
        logger.info("Found Turkish characters, using Turkish language")
        return 'tur+eng'
    
    try:
        # Sample multiple parts of the text for more accurate detection
        samples = []
        text_len = len(text)
        
        # Take samples from beginning, middle, and end
        if text_len > 6000:
            samples.append(text[:2000])
            samples.append(text[text_len//2-1000:text_len//2+1000])
            samples.append(text[-2000:])
        else:
            samples.append(text[:min(2000, text_len)])
        
        # Detect language for each sample
        lang_counts = {}
        for sample in samples:
            try:
                lang = detect(sample)
                lang_counts[lang] = lang_counts.get(lang, 0) + 1
            except:
                continue
        
        # Find most common language
        if lang_counts:
            most_common_lang = max(lang_counts.items(), key=lambda x: x[1])[0]
        else:
            most_common_lang = 'en'
        
        # Map language detector codes to our codes
        # Lithuanian might be detected as 'lt' by langdetect
        if most_common_lang in ['lt', 'lv', 'et', 'pl']:
            logger.info(f"Detected Baltic/Eastern European language ({most_common_lang}), using Lithuanian")
            most_common_lang = 'lt'
        # langdetect might falsely identify Lithuanian as other European languages
        elif most_common_lang in ['pt', 'ca', 'ro', 'cs', 'sk', 'sl']:
            # Check for common Lithuanian word patterns
            if any(word in text.lower() for word in ['teism', 'lietuv', 'valstyb', 'teisė', 'įstatymas']):
                logger.info(f"Detected {most_common_lang} but found Lithuanian keywords, overriding to Lithuanian")
                most_common_lang = 'lt'
        
        # Map to tesseract code or default to English
        tesseract_lang = LANGUAGE_CODES.get(most_common_lang, 'eng')
        
        # If detected language needs multiple languages for OCR
        if most_common_lang == 'lt':
            # Lithuanian OCR often works better with English as fallback
            tesseract_lang = 'lit+eng'
        elif most_common_lang == 'tr':
            # Turkish OCR often works better with English as fallback
            tesseract_lang = 'tur+eng'
        
        logger.info(f"Detected language: {most_common_lang} (tesseract code: {tesseract_lang})")
        return tesseract_lang
    
    except LangDetectException:
        logger.warning("Language detection failed, checking for specific patterns")
        
        # Check for Lithuanian word patterns
        if any(word in text.lower() for word in ['teism', 'lietuv', 'valstyb', 'teisė', 'įstatymas']):
            logger.info("Found Lithuanian keywords after failed detection")
            return 'lit+eng'
        
        logger.warning("No specific language patterns found, using English as fallback")
        return 'eng'

class DoclingExtractor(TextExtractor):
    """Text extractor implementation with improved OCR capabilities for scanned documents"""

    def extract(self, file_path):
        """
        Extract text from documents with advanced OCR fallback for scanned documents
        
        Args:
            file_path: Path to document file
            
        Returns:
            str: Extracted text content
        """
        # Get file extension for format-specific handling
        ext = file_path.split(".")[-1].lower()
        logger.info(f"Extracting text from {file_path} (format: {ext})")
        
        # Handle plain text files first (simplest case)
        if ext == "txt":
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
                    logger.info(f"Successfully extracted {len(text)} characters from text file")
                    return text
            except UnicodeDecodeError:
                # Try different encoding if UTF-8 fails
                try:
                    with open(file_path, "r", encoding="latin-1") as f:
                        text = f.read()
                        logger.info(f"Successfully extracted {len(text)} characters using latin-1 encoding")
                        return text
                except Exception as e:
                    logger.error(f"Failed to read text file with latin-1 encoding: {str(e)}")
                    return ""
            except Exception as e:
                logger.error(f"Failed to read text file: {str(e)}")
                return ""

        # PDF extraction with multiple fallback options
        if ext == "pdf":
            # Try docling first if available
            try:
                logger.info("Attempting to extract PDF with docling")
                # Import here to avoid errors if docling is not installed
                from docling.document_converter import DocumentConverter, PdfFormatOption
                from docling.datamodel.base_models import InputFormat
                from docling.datamodel.pipeline_options import PdfPipelineOptions
                from docling.models.tesseract_ocr_cli_model import TesseractCliOcrOptions

                # Configure docling options with high quality OCR
                ocr_options = TesseractCliOcrOptions()
                pdf_pipeline_options = PdfPipelineOptions(do_ocr=True, ocr_options=ocr_options, dpi=400)
                
                converter_options = {}
                converter_options[InputFormat.PDF] = PdfFormatOption(pipeline_options=pdf_pipeline_options)
                
                # Convert and extract text
                converter = DocumentConverter(format_options=converter_options)
                doc = converter.convert(file_path)
                text = doc.document.export_to_text()
                
                # Verify we extracted meaningful text
                if text and len(text.strip()) > 100:
                    logger.info(f"Successfully extracted {len(text)} characters with docling")
                    return text
                else:
                    logger.warning(f"Docling extracted insufficient text ({len(text)} chars), trying alternative methods")
                    raise ValueError("Insufficient text extracted")
                
            except ImportError as e:
                logger.warning(f"Docling not available, falling back to alternative methods: {e}")
                
                # Attempt to install docling if it's missing
                try:
                    logger.info("Attempting to install docling package...")
                    import subprocess
                    result = subprocess.run(['pip', 'install', 'docling'], 
                                           capture_output=True, text=True, check=False)
                    
                    if result.returncode == 0:
                        logger.info("Successfully installed docling, will be available for future runs")
                    else:
                        logger.warning(f"Failed to install docling: {result.stderr}")
                except Exception as install_error:
                    logger.warning(f"Error trying to install docling: {str(install_error)}")
            except Exception as e:
                logger.warning(f"Docling PDF extraction failed: {str(e)}")

            # Try PyMuPDF (fitz) - often better than PyPDF2 and pdfminer
            try:
                logger.info("Attempting to extract PDF with PyMuPDF")
                import fitz
                
                text = ""
                doc = fitz.open(file_path)
                
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    page_text = page.get_text()
                    text += page_text + "\n\n"
                
                if text and len(text.strip()) > 100:
                    logger.info(f"Successfully extracted {len(text)} characters with PyMuPDF")
                    return text
                else:
                    logger.warning(f"PyMuPDF extracted insufficient text ({len(text)} chars), trying pdfminer")
            except Exception as fitz_error:
                logger.warning(f"PDF extraction with PyMuPDF failed: {str(fitz_error)}")
                
            # Try pdfminer
            try:
                logger.info("Attempting to extract PDF with pdfminer")
                import pdfminer.high_level
                with open(file_path, 'rb') as file:
                    text = pdfminer.high_level.extract_text(file)
                    
                if text and len(text.strip()) > 100:
                    logger.info(f"Successfully extracted {len(text)} characters with pdfminer")
                    return text
                else:
                    logger.warning(f"PDFMiner extracted insufficient text ({len(text)} chars), trying PyPDF2")
            except Exception as pdf_error:
                logger.warning(f"PDF extraction with pdfminer failed: {str(pdf_error)}")
                    
            # Try PyPDF2
            try:
                logger.info("Attempting to extract PDF with PyPDF2")
                import PyPDF2
                text = ""
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    for page_num in range(len(reader.pages)):
                        page = reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"
                
                if text and len(text.strip()) > 100:
                    logger.info(f"Successfully extracted {len(text)} characters with PyPDF2")
                    return text
                else:
                    logger.warning(f"PyPDF2 extracted insufficient text ({len(text)} chars), trying OCR")
            except Exception as pypdf_error:
                logger.warning(f"PDF extraction with PyPDF2 failed: {str(pypdf_error)}")
                
            # If all standard methods failed, this is likely a scanned document - try OCR
            
            # Try direct Tesseract OCR via subprocess for best quality
            try:
                logger.info("Attempting to extract PDF with tesseract OCR directly")
                
                # Create a temporary directory for the images
                temp_dir = tempfile.mkdtemp()
                base_filename = os.path.join(temp_dir, "page")
                
                # Convert PDF to images using pdftoppm (comes with poppler-utils)
                # High quality conversion for better OCR results (300 DPI)
                subprocess.run(['pdftoppm', '-png', '-r', '300', file_path, base_filename], check=True)
                
                # Find all generated image files
                image_files = sorted(glob.glob(f"{base_filename}*.png"))
                
                if not image_files:
                    logger.warning("No images generated from PDF for OCR")
                    raise ValueError("PDF to image conversion failed")
                
                # First attempt with English to get initial text for language detection
                initial_text = ""
                for img_file in image_files[:1]:  # Just use first page for language detection
                    result = subprocess.run(
                        ['tesseract', img_file, 'stdout', '-l', 'eng', '--oem', '1', '--psm', '6'],
                        capture_output=True, text=True, check=True
                    )
                    initial_text += result.stdout
                
                # Detect language from the initial text
                lang_code = detect_language(initial_text)
                logger.info(f"OCR will use language: {lang_code}")
                
                # Now process all images with the detected language
                text = ""
                for img_file in image_files:
                    # Use high quality OCR settings with the detected language
                    result = subprocess.run(
                        ['tesseract', img_file, 'stdout', '-l', lang_code, '--oem', '1', '--psm', '6'],
                        capture_output=True, text=True, check=True
                    )
                    text += result.stdout + "\n\n"
                
                # Clean up temporary files
                for img_file in image_files:
                    os.remove(img_file)
                os.rmdir(temp_dir)
                
                if text and len(text.strip()) > 100:
                    logger.info(f"Successfully extracted {len(text)} characters with direct tesseract OCR")
                    return text
                else:
                    logger.warning(f"Tesseract OCR extracted insufficient text ({len(text)} chars), trying EasyOCR")
            except Exception as tess_error:
                logger.warning(f"PDF extraction with direct tesseract failed: {str(tess_error)}")
                
            # Try EasyOCR as a last resort (can be better for some document types)
            try:
                logger.info("Attempting to extract PDF with EasyOCR")
                import fitz  # PyMuPDF
                from PIL import Image
                import io
                import numpy as np
                import easyocr
                
                reader = easyocr.Reader(['en'])  # Initialize reader with English language
                
                doc = fitz.open(file_path)
                text = ""
                
                # Process each page
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    # Higher resolution for better OCR results
                    pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72))
                    img = Image.open(io.BytesIO(pix.tobytes()))
                    img_np = np.array(img)
                    
                    # Extract text using EasyOCR
                    results = reader.readtext(img_np)
                    page_text = " ".join([result[1] for result in results])
                    text += page_text + "\n\n"
                
                if text and len(text.strip()) > 100:
                    logger.info(f"Successfully extracted {len(text)} characters with EasyOCR")
                    return text
                else:
                    logger.error(f"All PDF extraction methods failed to extract meaningful text")
                    return ""
            except Exception as ocr_error:
                logger.error(f"PDF extraction with EasyOCR failed: {str(ocr_error)}")
                return ""
        
        # DOCX extraction
        elif ext == "docx":
            try:
                logger.info("Extracting DOCX with docx2txt")
                import docx2txt
                text = docx2txt.process(file_path)
                if text and len(text.strip()) > 0:
                    logger.info(f"Successfully extracted {len(text)} characters from DOCX")
                    return text
                else:
                    logger.warning("docx2txt extracted empty text, trying python-docx")
            except Exception as docx_error:
                logger.warning(f"DOCX extraction with docx2txt failed: {str(docx_error)}")
                
            # Try python-docx as fallback
            try:
                logger.info("Extracting DOCX with python-docx")
                from docx import Document
                doc = Document(file_path)
                
                # Extract paragraphs
                paragraphs = [paragraph.text for paragraph in doc.paragraphs]
                
                # Extract tables
                tables_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text for cell in row.cells]
                        tables_text.append(" | ".join(row_text))
                
                # Combine all text
                text = "\n".join(paragraphs) + "\n\n" + "\n".join(tables_text)
                
                if text and len(text.strip()) > 0:
                    logger.info(f"Successfully extracted {len(text)} characters from DOCX with python-docx")
                    return text
                else:
                    logger.error("Failed to extract text from DOCX with python-docx")
                    return ""
            except Exception as docx2_error:
                logger.error(f"DOCX extraction with python-docx failed: {str(docx2_error)}")
                return ""
        
        # DOC file handling (older Word format)
        elif ext == "doc":
            # Check if the filename indicates Lithuanian content
            filename = os.path.basename(file_path).lower()
            is_lithuanian_by_name = any(word in filename for word in ['teism', 'lietuv', 'valstyb', 'nutart'])
            if is_lithuanian_by_name:
                logger.info(f"DOC file name '{filename}' suggests Lithuanian content")
                
            # Try using Python-based antiword first (most reliable)
            try:
                logger.info("Extracting DOC with python-antiword")
                from antiword import Document
                doc = Document(file_path)
                text = doc.read()
                
                if text and len(text.strip()) > 100:
                    # Detect language
                    is_lithuanian_by_content = False
                    # Check for Lithuanian characters
                    if len(re.findall(r'[ąčęėįšųūž]', text[:5000])) > 2:
                        is_lithuanian_by_content = True
                        logger.info("Found Lithuanian characters in DOC content using python-antiword")
                    # Check for Lithuanian keywords
                    elif any(word in text.lower() for word in ['teism', 'lietuv', 'valstyb', 'teisė', 'įstatymas']):
                        is_lithuanian_by_content = True
                        logger.info("Found Lithuanian keywords in DOC content using python-antiword")
                    
                    # Determine language for cleaning
                    lang_code = detect_language(text[:1000])
                    if is_lithuanian_by_name or is_lithuanian_by_content:
                        logger.info("Overriding language detection to Lithuanian for DOC file")
                        lang_code = 'lit+eng'
                    
                    logger.info(f"DOC text language detected as: {lang_code}")
                    
                    # Clean up text to retain only meaningful content
                    cleaned_text = re.sub(r'\s+', ' ', text)  # normalize whitespace
                    cleaned_text = re.sub(r'[^\x20-\x7E\u0100-\u017F\u00C7\u00E7\u011E\u011F\u0130\u0131\u015E\u015F\u00D6\u00F6\u00DC\u00FC]', ' ', cleaned_text)  # remove non-printable and non-language chars
                    cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()  # clean up extra spaces
                    
                    logger.info(f"Successfully extracted {len(cleaned_text)} characters from DOC with python-antiword")
                    return cleaned_text
            except Exception as e:
                logger.warning(f"Python-antiword extraction failed: {str(e)}")
                
            # Try docx2txt as fallback (sometimes works for DOC files too)
            try:
                logger.info("Extracting DOC with docx2txt")
                import docx2txt
                
                # Create a temporary file with .docx extension
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                    temp_path = temp_file.name
                    # Copy the content to the temp file
                    with open(file_path, 'rb') as src_file:
                        temp_file.write(src_file.read())
                
                try:
                    # Try to process as docx
                    text = docx2txt.process(temp_path)
                    os.unlink(temp_path)  # Remove temp file
                    
                    if text and len(text.strip()) > 100:
                        # Detect language for cleaning
                        lang_code = detect_language(text[:1000])
                        if is_lithuanian_by_name:
                            logger.info("Overriding language detection to Lithuanian for DOC file")
                            lang_code = 'lit+eng'
                        
                        logger.info(f"DOC text language detected as: {lang_code}")
                        logger.info(f"Successfully extracted {len(text)} characters from DOC with docx2txt")
                        return text
                    else:
                        logger.warning("docx2txt extracted insufficient text, trying next method")
                except:
                    # Clean up temp file if extraction failed
                    if os.path.exists(temp_path):
                        os.unlink(temp_path)
                    raise
            except Exception as e:
                logger.warning(f"docx2txt extraction failed: {str(e)}")
            
            # Try system antiword if available (usually installed on Linux)
            try:
                logger.info("Extracting DOC with system antiword")
                # Ensure we're using the correct subprocess module
                import subprocess as sp
                result = sp.run(['antiword', file_path], capture_output=True, text=True)
                if result.returncode == 0 and result.stdout:
                    text = result.stdout
                    
                    # Check if we got meaningful text
                    if text and len(text.strip()) > 100:
                        # Clean up the text
                        cleaned_text = re.sub(r'\s+', ' ', text)  # normalize whitespace
                        cleaned_text = re.sub(r'[^\x20-\x7E\u0100-\u017F\u00C7\u00E7\u011E\u011F\u0130\u0131\u015E\u015F\u00D6\u00F6\u00DC\u00FC]', ' ', cleaned_text)  # remove non-printable chars
                        cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()  # clean up extra spaces
                        
                        # Detect language
                        if is_lithuanian_by_name:
                            logger.info("DOC is Lithuanian based on filename")
                            lang_code = 'lit+eng'
                        else:
                            lang_code = detect_language(cleaned_text[:1000])
                        
                        logger.info(f"DOC text language detected from antiword as: {lang_code}")
                        logger.info(f"Successfully extracted {len(cleaned_text)} characters from DOC with system antiword")
                        return cleaned_text
            except Exception as e:
                logger.warning(f"System antiword extraction failed: {str(e)}")
                
            # Try using catdoc if available (another DOC extractor)
            try:
                logger.info("Extracting DOC with catdoc")
                # Ensure we're using the correct subprocess module
                import subprocess as sp
                result = sp.run(['catdoc', '-d', 'utf-8', file_path], capture_output=True, text=True)
                if result.returncode == 0 and result.stdout:
                    text = result.stdout
                    
                    # Check if we got meaningful text
                    if text and len(text.strip()) > 100:
                        # Clean up the text
                        cleaned_text = re.sub(r'\s+', ' ', text)  # normalize whitespace
                        cleaned_text = re.sub(r'[^\x20-\x7E\u0100-\u017F\u00C7\u00E7\u011E\u011F\u0130\u0131\u015E\u015F\u00D6\u00F6\u00DC\u00FC]', ' ', cleaned_text)  # remove non-printable chars
                        cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()  # clean up extra spaces
                        
                        logger.info(f"Successfully extracted {len(cleaned_text)} characters from DOC with catdoc")
                        return cleaned_text
            except Exception as e:
                logger.warning(f"catdoc extraction failed: {str(e)}")
                
            # Try Apache Tika (very reliable for document extraction, especially DOC files)
            try:
                logger.info("Extracting DOC with Apache Tika")
                import tika
                from tika import parser
                
                # Initialize Tika
                tika.initVM()
                
                # Parse the document
                parsed = parser.from_file(file_path)
                if parsed and 'content' in parsed and parsed['content']:
                    text = parsed['content']
                    
                    # Clean up the text
                    cleaned_text = re.sub(r'\s+', ' ', text)  # normalize whitespace
                    cleaned_text = re.sub(r'[^\x20-\x7E\u0100-\u017F\u00C7\u00E7\u011E\u011F\u0130\u0131\u015E\u015F\u00D6\u00F6\u00DC\u00FC]', ' ', cleaned_text)  # remove non-printable chars
                    cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()  # clean up extra spaces
                    
                    if cleaned_text and len(cleaned_text.strip()) > 100:
                        # Detect language
                        lang_code = detect_language(cleaned_text[:1000])
                        if is_lithuanian_by_name:
                            logger.info("DOC is Lithuanian based on filename")
                            lang_code = 'lit+eng'
                        
                        logger.info(f"DOC text language detected from Tika as: {lang_code}")
                        logger.info(f"Successfully extracted {len(cleaned_text)} characters from DOC with Tika")
                        return cleaned_text
                    else:
                        logger.warning("Tika extracted insufficient text, trying next method")
            except Exception as e:
                logger.warning(f"Tika extraction failed: {str(e)}")
                
            # As a fallback, try a modified binary approach with enhanced filtering
            try:
                logger.info("Extracting DOC with enhanced binary extraction")
                with open(file_path, 'rb') as f:
                    content = f.read()
                
                # Extract 8-bit clean text strings (common in MS Word docs)
                strings = []
                current_string = []
                for i in range(len(content)):
                    c = content[i]
                    if (32 <= c <= 126) or (c >= 128):  # ASCII printable or extended ASCII
                        current_string.append(chr(c))
                    else:
                        if current_string:
                            s = ''.join(current_string)
                            if len(s) >= 4:  # only keep strings with at least 4 chars
                                strings.append(s)
                            current_string = []
                
                if current_string:
                    s = ''.join(current_string)
                    if len(s) >= 4:
                        strings.append(s)
                
                # Filter out strings that are likely not text
                filtered_strings = []
                for s in strings:
                    # Only keep strings that look like text
                    # Higher threshold for letter percentage
                    letter_count = sum(1 for c in s if c.isalpha())
                    if letter_count > 0 and letter_count / len(s) > 0.6:
                        # Further filter for strings that might have word-like patterns
                        if re.search(r'\b[a-zA-Z]{2,}\b', s):
                            filtered_strings.append(s)
                
                # Further filter by checking actual words
                word_like_strings = []
                for s in filtered_strings:
                    words = s.split()
                    # Keep only if it has a good ratio of real words (length > 1)
                    if words and sum(len(w) > 1 for w in words) / len(words) > 0.7:
                        word_like_strings.append(s)
                
                # Join filtered strings
                text = "\n".join(word_like_strings)
                
                # Final cleanup to remove gibberish and binary data
                text = re.sub(r'[^\x20-\x7E\u0100-\u017F\u00C7\u00E7\u011E\u011F\u0130\u0131\u015E\u015F\u00D6\u00F6\u00DC\u00FC]', ' ', text)
                text = re.sub(r'\s+', ' ', text).strip()
                
                if text and len(text.strip()) > 100:
                    logger.info(f"Successfully extracted {len(text)} characters from DOC with enhanced binary extraction")
                    return text
                else:
                    logger.warning("Enhanced binary extraction didn't yield useful text")
            except Exception as e:
                logger.warning(f"Enhanced binary extraction failed: {str(e)}")
            
            # Create fallback text
            fallback_text = f"""
DOCUMENT CONTENT EXTRACTION:
Filename: {os.path.basename(file_path)}
Size: {os.path.getsize(file_path)} bytes
Type: Microsoft Word Document (.doc)

This document could not be fully extracted due to format limitations.
Please convert the document to a more accessible format like PDF or DOCX.

RAGLens processed this document at: {os.path.basename(file_path)}
            """
            
            logger.info(f"Using fallback text for DOC file: {file_path}")
            return fallback_text
        
        # If all extraction methods fail, return empty string
        logger.error(f"No successful extraction method for {file_path}")
        return ""
